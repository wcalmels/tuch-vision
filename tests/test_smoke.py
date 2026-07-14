import json
from pathlib import Path

from docx import Document

from tuch_vision import __version__
from tuch_vision.profile import init_manuscript, load_profile
from tuch_vision.toc import (
    extract_headings,
    find_contents_span,
    run_toc,
)
from tuch_vision.vision.checks import run_checks
from tuch_vision.vision.docx_figures import FigureUnit


def test_version():
    assert __version__ == "0.2.1"


def test_init_and_profile(tmp_path: Path):
    root = tmp_path / "book"
    init_manuscript(root, name="Demo Book", ocr_hints=["alpha"])
    prof = load_profile(root)
    assert prof.name == "Demo Book"
    assert "alpha" in prof.ocr_hints
    assert (root / "state" / "manuscript_outline.json").exists()
    outline = json.loads((root / "state" / "manuscript_outline.json").read_text(encoding="utf-8"))
    assert outline["title"] == "Demo Book"


def test_checks_without_tth_hints():
    # conceptual unit should not require any theory priors
    u = FigureUnit(
        figure_id="c1",
        kind="conceptual",
        caption="A conceptual diagram",
    )
    out = run_checks(u, do_ocr=False, ocr_hints=None)
    assert out.status == "ok"


def _mini_book(path: Path) -> Path:
    doc = Document()
    doc.add_heading("Contents", level=1)
    doc.add_paragraph("Preface")
    doc.add_paragraph("Chapter 1: Alpha")
    doc.add_paragraph("Chapter 2: Beta")
    doc.add_heading("Preface", level=1)
    doc.add_paragraph("Once upon a time.")
    doc.add_heading("Chapter 1: Alpha", level=1)
    doc.add_paragraph("Body one.")
    doc.add_heading("Chapter 2: Beta", level=1)
    doc.add_paragraph("Body two.")
    doc.save(path)
    return path


def test_toc_audit_and_replace(tmp_path: Path):
    src = _mini_book(tmp_path / "mini.docx")
    report = run_toc(src, action="audit", max_level=1)
    assert report.heading_count >= 3
    assert find_contents_span(Document(str(src))) is not None
    assert any(h.text == "Chapter 1: Alpha" for h in extract_headings(Document(str(src))))

    out = tmp_path / "mini_toc.docx"
    report2 = run_toc(src, action="replace", max_level=2, output_docx=out)
    assert out.exists()
    assert report2.has_toc_field or "TOC field" in " ".join(i.message for i in report2.issues)

    # field markers in package
    from zipfile import ZipFile

    xml = ZipFile(out).read("word/document.xml").decode("utf-8", errors="replace")
    assert "TOC" in xml and "instrText" in xml


def test_toc_rebuild_static(tmp_path: Path):
    src = _mini_book(tmp_path / "mini2.docx")
    out = tmp_path / "mini_static.docx"
    report = run_toc(src, action="rebuild-static", max_level=1, output_docx=out)
    assert out.exists()
    doc = Document(str(out))
    span = find_contents_span(doc)
    assert span is not None
    body = [doc.paragraphs[i].text.strip() for i in range(span[0] + 1, span[1])]
    body = [t for t in body if t]
    assert "Chapter 1: Alpha" in body
    assert "Chapter 2: Beta" in body
    assert any("static TOC" in i.message for i in report.issues)
