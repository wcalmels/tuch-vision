"""Table of contents audit / Word TOC field / static rebuild."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

CONTENTS_TITLES = frozenset(
    {
        "contents",
        "table of contents",
        "índice",
        "indice",
        "tabla de contenidos",
    }
)


@dataclass
class HeadingEntry:
    index: int
    level: int
    text: str


@dataclass
class TocIssue:
    kind: str  # missing_in_toc | extra_in_toc | order | note
    message: str


@dataclass
class TocReport:
    source: str
    heading_count: int
    manual_toc_count: int
    has_toc_field: bool
    contents_span: tuple[int, int] | None
    issues: list[TocIssue] = field(default_factory=list)
    headings: list[HeadingEntry] = field(default_factory=list)
    manual_entries: list[str] = field(default_factory=list)
    action: str = "audit"

    @property
    def ok(self) -> bool:
        return not any(i.kind in {"missing_in_toc", "extra_in_toc", "order"} for i in self.issues)

    @property
    def summary(self) -> str:
        n_err = sum(1 for i in self.issues if i.kind != "note")
        n_note = sum(1 for i in self.issues if i.kind == "note")
        status = "OK" if self.ok else "MISMATCH"
        return (
            f"TOC {status}: {self.heading_count} headings (H1–H max), "
            f"{self.manual_toc_count} manual lines, field={self.has_toc_field}, "
            f"{n_err} issues, {n_note} notes"
        )

    def to_markdown(self) -> str:
        lines = [
            "# TOC Audit",
            "",
            f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')}",
            f"Source: `{self.source}`",
            f"Action: `{self.action}`",
            "",
            f"**{self.summary}**",
            "",
            "## Headings used for TOC",
            "",
        ]
        for h in self.headings:
            indent = "  " * (h.level - 1)
            lines.append(f"{indent}- H{h.level}: {h.text}")
        lines += ["", "## Manual TOC lines", ""]
        if not self.manual_entries:
            lines.append("_None found (or already field-based)._")
        else:
            for e in self.manual_entries:
                lines.append(f"- {e}")
        lines += ["", "## Issues", ""]
        if not self.issues:
            lines.append("_None._")
        else:
            for iss in self.issues:
                lines.append(f"- **{iss.kind}**: {iss.message}")
        lines += [
            "",
            "## Word tip",
            "",
            "After inserting a TOC field, open the DOCX in Word and press "
            "`Ctrl+A` then `F9` (or right-click TOC → Update Field) so page "
            "numbers resolve.",
            "",
        ]
        return "\n".join(lines)


def _style_name(p: Paragraph) -> str:
    try:
        return (p.style.name or "") if p.style else ""
    except Exception:
        return ""


def _heading_level(p: Paragraph) -> int | None:
    st = _style_name(p)
    m = re.match(r"Heading\s+(\d+)\s*$", st, re.I)
    if m:
        return int(m.group(1))
    return None


def is_contents_title(text: str) -> bool:
    return text.strip().lower() in CONTENTS_TITLES


def extract_headings(
    doc: Document,
    *,
    max_level: int = 2,
    skip_contents: bool = True,
) -> list[HeadingEntry]:
    out: list[HeadingEntry] = []
    for i, p in enumerate(doc.paragraphs):
        level = _heading_level(p)
        if level is None or level > max_level:
            continue
        text = p.text.strip()
        if not text:
            continue
        if skip_contents and is_contents_title(text):
            continue
        out.append(HeadingEntry(index=i, level=level, text=text))
    return out


def find_contents_span(doc: Document) -> tuple[int, int] | None:
    """Return [start, end) covering Contents heading + body until next Heading 1."""
    start: int | None = None
    for i, p in enumerate(doc.paragraphs):
        if is_contents_title(p.text) and (
            _heading_level(p) == 1 or _style_name(p).lower() in {"title", "normal"}
        ):
            # Prefer Heading 1 Contents; fall back to first Contents title
            if _heading_level(p) == 1 or start is None:
                start = i
                if _heading_level(p) == 1:
                    break
    if start is None:
        return None
    end = len(doc.paragraphs)
    for j in range(start + 1, len(doc.paragraphs)):
        if _heading_level(doc.paragraphs[j]) == 1:
            end = j
            break
    return start, end


def parse_manual_toc(doc: Document, span: tuple[int, int] | None) -> list[str]:
    if span is None:
        return []
    start, end = span
    entries: list[str] = []
    for i in range(start + 1, end):
        t = doc.paragraphs[i].text.strip()
        if not t:
            continue
        # skip leaders-only / page-number debris
        if re.fullmatch(r"[\d\.\s\-–—\.…]+", t):
            continue
        entries.append(t)
    return entries


def document_has_toc_field(docx_path: Path) -> bool:
    from zipfile import ZipFile

    with ZipFile(docx_path) as z:
        xml = z.read("word/document.xml").decode("utf-8", errors="replace")
    return bool(re.search(r"\bTOC\b", xml) and "instrText" in xml)


def _norm_title(text: str) -> str:
    t = text.strip().lower()
    t = re.sub(r"\s+", " ", t)
    t = re.sub(r"[.]+$", "", t)
    return t


def audit_headings_vs_manual(
    headings: list[HeadingEntry],
    manual: list[str],
    *,
    level: int = 1,
) -> list[TocIssue]:
    """Compare Heading-N titles to manual TOC lines (level 1 by default)."""
    issues: list[TocIssue] = []
    want = [_norm_title(h.text) for h in headings if h.level == level]
    have = [_norm_title(x) for x in manual]

    # Prefer matching H1 only against manual lines that look like chapter/part titles
    have_set = set(have)
    want_set = set(want)

    for title in want:
        if title not in have_set:
            issues.append(
                TocIssue("missing_in_toc", f"Heading not in manual TOC: {title}")
            )
    for title in have:
        if title not in want_set:
            # Manual often includes short leaders ("— Interludio —") — note, not hard fail
            if title.startswith("—") or title.startswith("-"):
                issues.append(TocIssue("note", f"Decorative TOC line: {title}"))
            else:
                issues.append(
                    TocIssue("extra_in_toc", f"Manual TOC line not a Heading {level}: {title}")
                )

    # Order: compare filtered intersection sequence
    want_in_have = [t for t in want if t in have_set]
    have_in_want = [t for t in have if t in want_set]
    if want_in_have != have_in_want:
        issues.append(
            TocIssue(
                "order",
                "Manual TOC order differs from document Heading order "
                f"(first drift near {want_in_have[:3]} vs {have_in_want[:3]})",
            )
        )
    return issues


def _remove_paragraph(p: Paragraph) -> None:
    el = p._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def _insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _add_toc_field(paragraph: Paragraph, *, levels: str = "1-2") -> None:
    """Insert a Word TOC field (page numbers resolve after Word Update Field)."""
    run = paragraph.add_run()
    r = run._r

    fd_begin = OxmlElement("w:fldChar")
    fd_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f'TOC \\o "{levels}" \\h \\z \\u'

    fd_sep = OxmlElement("w:fldChar")
    fd_sep.set(qn("w:fldCharType"), "separate")

    # Placeholder until Word refreshes
    text = OxmlElement("w:t")
    text.text = (
        "Right-click → Update Field (or Ctrl+A, F9) to build the table of contents."
    )

    fd_end = OxmlElement("w:fldChar")
    fd_end.set(qn("w:fldCharType"), "end")

    r.append(fd_begin)
    r.append(instr)
    r.append(fd_sep)
    r.append(text)
    r.append(fd_end)


def demote_contents_heading(doc: Document, idx: int) -> None:
    """Avoid self-listing: Contents should not be Heading 1 for Word TOC fields."""
    p = doc.paragraphs[idx]
    try:
        p.style = "Title"
    except KeyError:
        try:
            p.style = "Normal"
        except KeyError:
            pass


def replace_manual_with_field(
    doc: Document,
    *,
    max_level: int = 2,
) -> tuple[int, int]:
    """
    Remove manual TOC body under Contents and insert a TOC field.
    Returns (removed_paras, start_index).
    """
    span = find_contents_span(doc)
    if span is None:
        # Create Contents + field at top of body (before first Heading 1)
        first_h1 = next(
            (i for i, p in enumerate(doc.paragraphs) if _heading_level(p) == 1),
            0,
        )
        anchor = doc.paragraphs[max(0, first_h1)]
        # insert before first H1
        new_p = OxmlElement("w:p")
        anchor._element.addprevious(new_p)
        title = Paragraph(new_p, anchor._parent)
        try:
            title.style = "Title"
        except KeyError:
            pass
        title.add_run("Contents")
        field_p = _insert_paragraph_after(title)
        _add_toc_field(field_p, levels=f"1-{max_level}")
        return 0, 0

    start, end = span
    demote_contents_heading(doc, start)
    # Delete body paras (not the title)
    removed = 0
    for idx in range(end - 1, start, -1):
        _remove_paragraph(doc.paragraphs[idx])
        removed += 1
    # Re-resolve title paragraph after deletes
    title_p = doc.paragraphs[start]
    field_p = _insert_paragraph_after(title_p)
    _add_toc_field(field_p, levels=f"1-{max_level}")
    return removed, start


def rebuild_static_toc(
    doc: Document,
    *,
    max_level: int = 2,
    include_levels: Iterable[int] | None = None,
) -> int:
    """Replace manual TOC body with a static list from current headings."""
    levels = set(include_levels or range(1, max_level + 1))
    headings = extract_headings(doc, max_level=max_level, skip_contents=True)
    span = find_contents_span(doc)
    if span is None:
        raise RuntimeError("No Contents section found to rebuild")
    start, end = span
    for idx in range(end - 1, start, -1):
        _remove_paragraph(doc.paragraphs[idx])
    title_p = doc.paragraphs[start]
    anchor = title_p
    n = 0
    for h in headings:
        if h.level not in levels:
            continue
        new_p = _insert_paragraph_after(anchor)
        try:
            new_p.style = "Normal"
        except KeyError:
            pass
        prefix = ("  " * (h.level - 1)) if h.level > 1 else ""
        new_p.add_run(f"{prefix}{h.text}")
        anchor = new_p
        n += 1
    return n


def run_toc(
    docx: Path,
    *,
    action: str = "audit",
    max_level: int = 2,
    output_docx: Path | None = None,
) -> TocReport:
    """
    action: audit | replace | rebuild-static
    """
    docx = Path(docx)
    doc = Document(str(docx))
    headings = extract_headings(doc, max_level=max_level, skip_contents=True)
    span = find_contents_span(doc)
    manual = parse_manual_toc(doc, span)
    has_field = document_has_toc_field(docx)

    issues = audit_headings_vs_manual(headings, manual, level=1)
    if has_field:
        issues.append(
            TocIssue(
                "note",
                "DOCX already contains a TOC field — open in Word and Update Field for pages.",
            )
        )
    if span is None:
        issues.append(TocIssue("note", "No Contents/Índice section detected."))

    report = TocReport(
        source=str(docx),
        heading_count=len(headings),
        manual_toc_count=len(manual),
        has_toc_field=has_field,
        contents_span=span,
        issues=issues,
        headings=headings,
        manual_entries=manual,
        action=action,
    )

    if action == "audit":
        return report

    if action == "replace":
        removed, _ = replace_manual_with_field(doc, max_level=max_level)
        report.issues.append(
            TocIssue("note", f"Replaced manual TOC body ({removed} paras) with Word TOC field.")
        )
        report.action = "replace"
    elif action == "rebuild-static":
        n = rebuild_static_toc(doc, max_level=max_level)
        report.issues.append(TocIssue("note", f"Rebuilt static TOC with {n} heading lines."))
        report.action = "rebuild-static"
    else:
        raise ValueError(f"Unknown action: {action}")

    out = Path(output_docx) if output_docx else docx.with_name(docx.stem + "_toc.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    report.source = f"{docx} → {out}"
    report.has_toc_field = action == "replace" or document_has_toc_field(out)
    # refresh manual counts after write for report clarity
    if action == "replace":
        report.manual_toc_count = 0
        report.manual_entries = []
    return report
